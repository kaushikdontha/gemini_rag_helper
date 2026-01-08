# document_processor.py - Multi-format document loader and text chunker

import os
import re
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import tiktoken

# Document loaders
from pypdf import PdfReader
from docx import Document as DocxDocument


@dataclass
class DocumentChunk:
    """Represents a chunk of document text with metadata."""
    content: str
    metadata: Dict[str, Any]
    chunk_id: int
    token_count: int


class DocumentProcessor:
    """
    Multi-format document processor that extracts text and creates chunks.
    Supports: PDF, DOCX, TXT, Markdown
    """
    
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.txt', '.md', '.markdown'}
    
    def __init__(self, chunk_size: int = 750, chunk_overlap: int = 100):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Target size of each chunk in tokens (500-1000 range)
            chunk_overlap: Number of overlapping tokens between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
    
    def is_supported(self, filename: str) -> bool:
        """Check if the file format is supported."""
        ext = os.path.splitext(filename)[1].lower()
        return ext in self.SUPPORTED_FORMATS
    
    def count_tokens(self, text: str) -> int:
        """Count the number of tokens in text."""
        return len(self.tokenizer.encode(text))
    
    def extract_text(self, file_content: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Extract text from document bytes.
        
        Args:
            file_content: Raw file content as bytes
            filename: Name of the file for format detection
            
        Returns:
            List of dicts with 'text' and 'metadata' (page/section info)
        """
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == '.pdf':
            return self._extract_pdf(file_content, filename)
        elif ext == '.docx':
            return self._extract_docx(file_content, filename)
        elif ext in {'.txt', '.md', '.markdown'}:
            return self._extract_text_file(file_content, filename, ext)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def _extract_pdf(self, file_content: bytes, filename: str) -> List[Dict[str, Any]]:
        """Extract text from PDF file."""
        import io
        pages = []
        pdf_reader = PdfReader(io.BytesIO(file_content))
        
        for page_num, page in enumerate(pdf_reader.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                pages.append({
                    'text': text.strip(),
                    'metadata': {
                        'document_name': filename,
                        'page': page_num,
                        'total_pages': len(pdf_reader.pages),
                        'source_type': 'pdf'
                    }
                })
        
        return pages
    
    def _extract_docx(self, file_content: bytes, filename: str) -> List[Dict[str, Any]]:
        """Extract text from DOCX file."""
        import io
        doc = DocxDocument(io.BytesIO(file_content))
        
        sections = []
        current_section = []
        current_heading = "Document Start"
        section_num = 1
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if this is a heading
            if para.style.name.startswith('Heading'):
                # Save previous section if exists
                if current_section:
                    sections.append({
                        'text': '\n'.join(current_section),
                        'metadata': {
                            'document_name': filename,
                            'section': section_num,
                            'section_title': current_heading,
                            'source_type': 'docx'
                        }
                    })
                    section_num += 1
                
                current_heading = text
                current_section = [text]
            else:
                current_section.append(text)
        
        # Don't forget the last section
        if current_section:
            sections.append({
                'text': '\n'.join(current_section),
                'metadata': {
                    'document_name': filename,
                    'section': section_num,
                    'section_title': current_heading,
                    'source_type': 'docx'
                }
            })
        
        return sections if sections else [{
            'text': '\n'.join([p.text for p in doc.paragraphs if p.text.strip()]),
            'metadata': {
                'document_name': filename,
                'section': 1,
                'source_type': 'docx'
            }
        }]
    
    def _extract_text_file(self, file_content: bytes, filename: str, ext: str) -> List[Dict[str, Any]]:
        """Extract text from TXT or Markdown file."""
        try:
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            text = file_content.decode('latin-1')
        
        source_type = 'markdown' if ext in {'.md', '.markdown'} else 'text'
        
        # For markdown, try to split by headers
        if source_type == 'markdown':
            sections = self._split_markdown_sections(text, filename)
            if sections:
                return sections
        
        # Default: return as single section
        return [{
            'text': text.strip(),
            'metadata': {
                'document_name': filename,
                'section': 1,
                'source_type': source_type
            }
        }]
    
    def _split_markdown_sections(self, text: str, filename: str) -> List[Dict[str, Any]]:
        """Split markdown text by headers."""
        # Pattern to match markdown headers
        header_pattern = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
        
        sections = []
        last_end = 0
        section_num = 1
        current_heading = "Document Start"
        
        for match in header_pattern.finditer(text):
            # Get content before this header
            content = text[last_end:match.start()].strip()
            if content:
                sections.append({
                    'text': content,
                    'metadata': {
                        'document_name': filename,
                        'section': section_num,
                        'section_title': current_heading,
                        'source_type': 'markdown'
                    }
                })
                section_num += 1
            
            current_heading = match.group(2)
            last_end = match.start()
        
        # Get remaining content
        if last_end < len(text):
            content = text[last_end:].strip()
            if content:
                sections.append({
                    'text': content,
                    'metadata': {
                        'document_name': filename,
                        'section': section_num,
                        'section_title': current_heading,
                        'source_type': 'markdown'
                    }
                })
        
        return sections
    
    def create_chunks(self, documents: List[Dict[str, Any]]) -> List[DocumentChunk]:
        """
        Create token-based chunks from extracted documents.
        
        Args:
            documents: List of extracted document sections
            
        Returns:
            List of DocumentChunk objects
        """
        all_chunks = []
        global_chunk_id = 0
        
        for doc in documents:
            text = doc['text']
            metadata = doc['metadata'].copy()
            
            # Split text into sentences for better chunking
            sentences = self._split_into_sentences(text)
            
            current_chunk = []
            current_tokens = 0
            
            for sentence in sentences:
                sentence_tokens = self.count_tokens(sentence)
                
                # If single sentence exceeds chunk size, split it
                if sentence_tokens > self.chunk_size:
                    # Save current chunk if exists
                    if current_chunk:
                        chunk_text = ' '.join(current_chunk)
                        all_chunks.append(DocumentChunk(
                            content=chunk_text,
                            metadata=metadata.copy(),
                            chunk_id=global_chunk_id,
                            token_count=self.count_tokens(chunk_text)
                        ))
                        global_chunk_id += 1
                        current_chunk = []
                        current_tokens = 0
                    
                    # Split long sentence into smaller parts
                    words = sentence.split()
                    temp_chunk = []
                    temp_tokens = 0
                    
                    for word in words:
                        word_tokens = self.count_tokens(word + ' ')
                        if temp_tokens + word_tokens > self.chunk_size:
                            if temp_chunk:
                                chunk_text = ' '.join(temp_chunk)
                                all_chunks.append(DocumentChunk(
                                    content=chunk_text,
                                    metadata=metadata.copy(),
                                    chunk_id=global_chunk_id,
                                    token_count=self.count_tokens(chunk_text)
                                ))
                                global_chunk_id += 1
                            temp_chunk = [word]
                            temp_tokens = word_tokens
                        else:
                            temp_chunk.append(word)
                            temp_tokens += word_tokens
                    
                    if temp_chunk:
                        current_chunk = temp_chunk
                        current_tokens = temp_tokens
                
                # Normal case: add sentence to current chunk
                elif current_tokens + sentence_tokens <= self.chunk_size:
                    current_chunk.append(sentence)
                    current_tokens += sentence_tokens
                else:
                    # Save current chunk and start new one
                    if current_chunk:
                        chunk_text = ' '.join(current_chunk)
                        all_chunks.append(DocumentChunk(
                            content=chunk_text,
                            metadata=metadata.copy(),
                            chunk_id=global_chunk_id,
                            token_count=self.count_tokens(chunk_text)
                        ))
                        global_chunk_id += 1
                        
                        # Create overlap from end of current chunk
                        overlap_text = self._get_overlap_text(current_chunk)
                        current_chunk = [overlap_text, sentence] if overlap_text else [sentence]
                        current_tokens = self.count_tokens(' '.join(current_chunk))
                    else:
                        current_chunk = [sentence]
                        current_tokens = sentence_tokens
            
            # Don't forget the last chunk
            if current_chunk:
                chunk_text = ' '.join(current_chunk)
                all_chunks.append(DocumentChunk(
                    content=chunk_text,
                    metadata=metadata.copy(),
                    chunk_id=global_chunk_id,
                    token_count=self.count_tokens(chunk_text)
                ))
                global_chunk_id += 1
        
        return all_chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences."""
        # Simple sentence splitting
        sentence_pattern = re.compile(r'(?<=[.!?])\s+(?=[A-Z])')
        sentences = sentence_pattern.split(text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _get_overlap_text(self, sentences: List[str]) -> str:
        """Get overlap text from the end of sentences list."""
        overlap_sentences = []
        overlap_tokens = 0
        
        for sentence in reversed(sentences):
            sentence_tokens = self.count_tokens(sentence)
            if overlap_tokens + sentence_tokens <= self.chunk_overlap:
                overlap_sentences.insert(0, sentence)
                overlap_tokens += sentence_tokens
            else:
                break
        
        return ' '.join(overlap_sentences)
    
    def process_document(self, file_content: bytes, filename: str) -> List[DocumentChunk]:
        """
        Complete processing pipeline: extract text and create chunks.
        
        Args:
            file_content: Raw file content
            filename: Name of the file
            
        Returns:
            List of DocumentChunk objects ready for embedding
        """
        if not self.is_supported(filename):
            raise ValueError(f"Unsupported file format: {filename}")
        
        # Extract text sections
        documents = self.extract_text(file_content, filename)
        
        # Create chunks
        chunks = self.create_chunks(documents)
        
        return chunks
